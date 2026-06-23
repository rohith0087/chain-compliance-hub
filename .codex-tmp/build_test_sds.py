from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("artifacts/Test_Supplier_SDS_Test_Document.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
RED = "9B1C1C"
WHITE = "FFFFFF"
MUTED = "666666"
BLACK = "000000"


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def write_cell(cell, text, bold=False, color=BLACK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    set_run_font(p.add_run(text), size=size, color=color, bold=bold)


def add_key_value_table(doc, rows, label_width=2500, value_width=6860):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        write_cell(cells[0], label, bold=True, color=DARK_BLUE)
        shade_cell(cells[0], LIGHT_GRAY)
        write_cell(cells[1], value)
    set_table_geometry(table, [label_width, value_width])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_data_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, header in enumerate(headers):
        write_cell(table.rows[0].cells[i], header, bold=True, color=WHITE, size=9)
        shade_cell(table.rows[0].cells[i], DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            write_cell(cells[i], value, size=9)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), bold=True)
        set_run_font(p.add_run(text[len(bold_lead):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_section_heading(doc, number, title):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(f"SECTION {number}: {title.upper()}"), size=16, color=BLUE, bold=True)
    return p


def add_field(field_paragraph, field_code):
    run = field_paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hp.paragraph_format.space_after = Pt(0)
set_run_font(hp.add_run("TEST FIXTURE - NOT FOR REGULATORY OR SAFETY USE"), size=9, color=RED, bold=True)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(0)
fp.paragraph_format.space_after = Pt(0)
set_run_font(fp.add_run("TR2C SDS Extraction Test | SDS-TS-NSC100-2026-001 | Page "), size=9, color=MUTED)
add_field(fp, "PAGE")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(3)
set_run_font(p.add_run("SAFETY DATA SHEET (SDS)"), size=24, color=DARK_BLUE, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(14)
set_run_font(p.add_run("TR2C Neutral Surface Cleaner NSC-100"), size=15, color=MUTED, bold=True)

notice = doc.add_table(rows=1, cols=1)
notice.style = "Table Grid"
write_cell(
    notice.cell(0, 0),
    "CONTROLLED TEST DOCUMENT - FICTIONAL DATA ONLY. Created solely to test document upload, extraction, verification, evidence sharing, expiration tracking, and audit workflows. Do not use for workplace safety, transport, emergency response, or regulatory compliance.",
    bold=True,
    color=RED,
    size=10,
)
shade_cell(notice.cell(0, 0), "FDECEC")
set_table_geometry(notice, [9360])
doc.add_paragraph().paragraph_format.space_after = Pt(0)

add_key_value_table(doc, [
    ("Submitting supplier", "Test Supplier"),
    ("Submitted to", "Test Buyer"),
    ("Product name", "TR2C Neutral Surface Cleaner NSC-100"),
    ("SDS number", "SDS-TS-NSC100-2026-001"),
    ("Certification number", "CERT-SDS-2026-0619-001"),
    ("Issue date", "June 19, 2026 (2026-06-19)"),
    ("Revision date", "June 19, 2026 (2026-06-19)"),
    ("Expiration date", "June 19, 2027 (2027-06-19)"),
    ("Document version", "1.0"),
    ("Intended standard", "OSHA Hazard Communication Standard, 29 CFR 1910.1200 - test fixture only"),
])

add_section_heading(doc, 1, "Identification")
add_key_value_table(doc, [
    ("Product identifier", "TR2C Neutral Surface Cleaner NSC-100"),
    ("Recommended use", "Fictional water-based hard-surface cleaner used only as software test data"),
    ("Restrictions on use", "Not a real commercial product; do not distribute or use"),
    ("Supplier", "Test Supplier, 100 Test Park Drive, Chicago, IL 60601, United States"),
    ("Supplier contact", "Amit | amit@tracer2c.com | +1 (555) 010-2026"),
    ("Emergency telephone", "+1 (555) 010-0911 - NONFUNCTIONAL TEST NUMBER"),
])

doc.add_page_break()
add_section_heading(doc, 2, "Hazard identification")
add_body(doc, "OSHA classification: Not classified as hazardous under the criteria represented in this fictional test document.", "OSHA classification:")
add_key_value_table(doc, [
    ("Signal word", "None required"),
    ("Hazard statements", "None assigned for this fictional mixture"),
    ("Precautionary statements", "Avoid eye contact. Use ordinary workplace hygiene. Keep out of reach of children."),
    ("Other hazards", "Spilled liquid may create a slip hazard."),
])

add_section_heading(doc, 3, "Composition / information on ingredients")
add_data_table(doc, ["Ingredient", "CAS number", "Concentration"], [
    ("Water", "7732-18-5", "85-95%"),
    ("Nonionic surfactant blend", "Proprietary", "3-5%"),
    ("Sodium citrate", "68-04-2", "1-3%"),
    ("Propylene glycol", "57-55-6", "1-3%"),
], [4200, 2400, 2760])
add_body(doc, "The concentrations above are fictional and are included only to test structured extraction from a composition table.")

add_section_heading(doc, 4, "First-aid measures")
add_key_value_table(doc, [
    ("Eye contact", "Rinse cautiously with water for several minutes. Obtain medical advice if irritation persists."),
    ("Skin contact", "Wash with soap and water. Remove contaminated clothing if needed."),
    ("Inhalation", "Move to fresh air if discomfort occurs."),
    ("Ingestion", "Rinse mouth. Do not induce vomiting. Contact a medical professional if unwell."),
    ("Most important symptoms", "Temporary eye irritation may occur."),
])

add_section_heading(doc, 5, "Fire-fighting measures")
add_body(doc, "Suitable extinguishing media: Use extinguishing media appropriate for the surrounding fire. The fictional water-based mixture is not expected to be flammable.", "Suitable extinguishing media:")
add_body(doc, "Protective equipment: Firefighters should use standard protective equipment appropriate to the surrounding incident.", "Protective equipment:")

doc.add_page_break()
add_section_heading(doc, 6, "Accidental release measures")
add_body(doc, "Isolate the area, use ordinary protective gloves, absorb with inert material, and clean the surface with water. Prevent entry into drains in large quantities. Mark wet floors to prevent slips.")

add_section_heading(doc, 7, "Handling and storage")
add_key_value_table(doc, [
    ("Safe handling", "Avoid eye contact and prolonged skin contact. Do not mix with other cleaning products."),
    ("Storage conditions", "Keep container closed in a cool, dry, ventilated place away from direct sunlight."),
    ("Storage temperature", "5-30 C (41-86 F)"),
    ("Incompatible materials", "Strong oxidizers and strong acids"),
])

add_section_heading(doc, 8, "Exposure controls / personal protection")
add_key_value_table(doc, [
    ("Exposure limits", "No occupational exposure limits assigned for the fictional mixture"),
    ("Engineering controls", "General ventilation is adequate for intended test representation"),
    ("Eye protection", "Safety glasses if splashing is possible"),
    ("Hand protection", "Protective gloves for prolonged contact"),
    ("Respiratory protection", "Not normally required under represented conditions"),
])

add_section_heading(doc, 9, "Physical and chemical properties")
add_data_table(doc, ["Property", "Test value", "Property", "Test value"], [
    ("Appearance", "Clear blue liquid", "Odor", "Mild citrus"),
    ("pH", "7.0-8.0", "Freezing point", "Approximately 0 C"),
    ("Boiling point", "Approximately 100 C", "Flash point", "Not applicable"),
    ("Relative density", "1.01 at 20 C", "Solubility", "Completely soluble in water"),
], [1900, 2780, 1900, 2780])

doc.add_page_break()
add_section_heading(doc, 10, "Stability and reactivity")
add_key_value_table(doc, [
    ("Reactivity", "No dangerous reaction expected under normal represented conditions"),
    ("Chemical stability", "Stable under recommended fictional storage conditions"),
    ("Conditions to avoid", "Freezing, excessive heat, and direct sunlight"),
    ("Incompatible materials", "Strong oxidizers and strong acids"),
    ("Hazardous decomposition", "None expected under normal represented use"),
])

add_section_heading(doc, 11, "Toxicological information")
add_body(doc, "Likely routes of exposure: Eye contact, skin contact, inhalation of mist, and ingestion.", "Likely routes of exposure:")
add_body(doc, "Expected effects: Temporary eye irritation and mild skin irritation after prolonged contact. No test data are available because this is a fictional software test mixture.", "Expected effects:")
add_body(doc, "Carcinogenicity: No components are intentionally represented as listed by IARC, NTP, or OSHA at reportable concentrations.", "Carcinogenicity:")

add_section_heading(doc, 12, "Ecological information")
add_body(doc, "Ecotoxicity, persistence, bioaccumulation, and mobility data have not been determined for this fictional mixture. Avoid uncontrolled release to the environment.")

add_section_heading(doc, 13, "Disposal considerations")
add_body(doc, "Dispose of contents and container in accordance with applicable local requirements. This statement is for test extraction only and is not disposal guidance for any actual product.")

add_section_heading(doc, 14, "Transport information")
add_key_value_table(doc, [
    ("UN number", "Not regulated for this fictional test record"),
    ("Proper shipping name", "Not applicable"),
    ("Transport hazard class", "Not applicable"),
    ("Packing group", "Not applicable"),
    ("Marine pollutant", "No"),
])

add_section_heading(doc, 15, "Regulatory information")
add_body(doc, "This fictional document is formatted to resemble the 16-section structure used for OSHA Hazard Communication SDS records. It is not a regulatory determination, certification, or legal opinion.")

add_section_heading(doc, 16, "Other information")
add_key_value_table(doc, [
    ("Prepared by", "Test Supplier Compliance Team"),
    ("Approved by", "Amit, Supplier Test Administrator"),
    ("Certification number", "CERT-SDS-2026-0619-001"),
    ("Expiration date", "June 19, 2027 (2027-06-19)"),
    ("Revision summary", "Initial fictional version created for TraceR2C end-to-end testing"),
])
add_body(doc, "Test signature: /s/ Amit - Test Supplier - 2026-06-19", "Test signature:")

final_notice = doc.add_table(rows=1, cols=1)
final_notice.style = "Table Grid"
write_cell(final_notice.cell(0, 0), "END OF CONTROLLED TEST FIXTURE - NOT VALID FOR REAL-WORLD USE", bold=True, color=RED, size=10)
shade_cell(final_notice.cell(0, 0), "FDECEC")
set_table_geometry(final_notice, [9360])

props = doc.core_properties
props.title = "Test Supplier SDS Test Document"
props.subject = "Fictional SDS fixture for TraceR2C extraction and evidence-sharing testing"
props.author = "TraceR2C Test Fixture Generator"
props.keywords = "TEST FIXTURE, SDS, Test Supplier, Test Buyer"
props.comments = "Fictional data only. Not for regulatory or safety use."

doc.save(OUT)
print(OUT.resolve())
